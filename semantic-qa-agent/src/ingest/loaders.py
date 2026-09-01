"""Document loaders -- PRD Section 9.1.

One loader per format, all returning the same `RawDocument` shape so the rest of
the pipeline never branches on file type.

Design decisions worth stating:

* **Per-page extraction for PDFs.** `page_no` is read from the parser, never
  estimated from character offsets. A citation that says "page 12" must be
  actually true, or the whole citation feature is worthless.
* **Optional heavy dependencies.** PyMuPDF and python-docx are imported lazily.
  A missing optional dependency degrades that one format to a clear, actionable
  error -- it never prevents the pipeline from running on the formats that do
  work (PRD principle #6: fail loudly, degrade gracefully).
* **Every failure is reported, never swallowed.** A file that cannot be parsed
  is recorded on the document with a `load_error` and surfaced in the run
  summary and the unparsed report.
"""

from __future__ import annotations

import csv
import html as html_lib
import io
import re
from pathlib import Path

from ..config import Config
from ..models import Page, RawDocument
from ..utils.logging import get_logger

_HTML_TAG = re.compile(r"<[^>]+>")
_HTML_DROP_BLOCK = re.compile(
    r"<(script|style|noscript)\b.*?</\1>", re.IGNORECASE | re.DOTALL
)


def make_doc_id(path: Path, root: Path) -> str:
    """Stable, readable, filesystem-independent document id."""
    try:
        rel = path.resolve().relative_to(root.resolve())
    except ValueError:
        rel = Path(path.name)
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", rel.as_posix().rsplit(".", 1)[0])
    return slug.strip("_").lower() or "doc"


def _title_from_path(path: Path) -> str:
    return path.stem.replace("_", " ").replace("-", " ").strip().title()


# --------------------------------------------------------------------------- #
# Per-format loaders
# --------------------------------------------------------------------------- #

def load_pdf(path: Path) -> tuple[list[Page], str | None]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return [], "PyMuPDF not installed -- run: pip install pymupdf"
    try:
        pages: list[Page] = []
        with fitz.open(path) as doc:
            if doc.is_encrypted and not doc.authenticate(""):
                return [], "encrypted PDF -- cannot extract text"
            for i, page in enumerate(doc, start=1):
                pages.append(Page(page_no=i, text=page.get_text("text") or ""))
        return pages, None
    except Exception as exc:  # a corrupt file must not stop the run
        return [], f"pdf parse error: {exc}"


def load_docx(path: Path) -> tuple[list[Page], str | None]:
    try:
        import docx  # python-docx
    except ImportError:
        return [], "python-docx not installed -- run: pip install python-docx"
    try:
        document = docx.Document(str(path))
        parts: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Keep heading structure as markdown so the chunker's heading
            # detection works uniformly across formats.
            style = (para.style.name or "").lower()
            if style.startswith("heading"):
                level = "".join(c for c in style if c.isdigit()) or "1"
                parts.append(f"{'#' * min(int(level), 6)} {text}")
            else:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return [Page(page_no=1, text="\n\n".join(parts))], None
    except Exception as exc:
        return [], f"docx parse error: {exc}"


def load_text(path: Path) -> tuple[list[Page], str | None]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [Page(page_no=1, text=text)], None
    except Exception as exc:
        return [], f"text read error: {exc}"


def load_html(path: Path) -> tuple[list[Page], str | None]:
    """Prefer BeautifulSoup; fall back to a regex strip so HTML always works."""
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return [], f"html read error: {exc}"
    try:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text("\n")
    except ImportError:
        raw = _HTML_DROP_BLOCK.sub(" ", raw)
        text = html_lib.unescape(_HTML_TAG.sub("\n", raw))
    return [Page(page_no=1, text=text)], None


def load_csv(path: Path) -> tuple[list[Page], str | None]:
    """Flatten rows to 'col: value' sentences.

    A raw CSV row ("A,3,true") embeds terribly -- the column names carry all the
    meaning. Restating each row as labelled fields makes rows genuinely
    retrievable by a semantic query.
    """
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        sample = raw[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.DictReader(io.StringIO(raw), dialect=dialect)
        lines: list[str] = []
        for i, row in enumerate(reader, start=1):
            fields = [
                f"{(k or '').strip()}: {(v or '').strip()}"
                for k, v in row.items()
                if v and str(v).strip()
            ]
            if fields:
                lines.append(f"Row {i} -- " + "; ".join(fields))
        return [Page(page_no=1, text="\n".join(lines))], None
    except Exception as exc:
        return [], f"csv parse error: {exc}"


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".md": load_text,
    ".txt": load_text,
    ".html": load_html,
    ".htm": load_html,
    ".csv": load_csv,
}


# --------------------------------------------------------------------------- #
# Discovery + dispatch
# --------------------------------------------------------------------------- #

def discover_files(root: Path, cfg: Config) -> tuple[list[Path], list[Path]]:
    """Return (supported, unsupported) files under `root`."""
    if not root.exists():
        raise FileNotFoundError(f"input directory does not exist: {root}")
    pattern = "**/*" if cfg.ingest.recursive else "*"
    supported: list[Path] = []
    unsupported: list[Path] = []
    for path in sorted(root.glob(pattern)):
        if not path.is_file() or path.name.startswith("."):
            continue
        (supported if path.suffix.lower() in cfg.ingest.supported_extensions
         else unsupported).append(path)
    return supported, unsupported


def load_document(path: Path, cfg: Config, root: Path) -> RawDocument:
    """Parse one file into a RawDocument, never raising on bad input."""
    log = get_logger()
    doc = RawDocument(
        doc_id=make_doc_id(path, root),
        source_path=str(path),
        doc_title=_title_from_path(path),
        doc_type=path.suffix.lower().lstrip("."),
        bytes_size=path.stat().st_size,
    )

    size_mb = doc.bytes_size / (1024 * 1024)
    if size_mb > cfg.ingest.max_file_mb:
        doc.load_error = f"file too large ({size_mb:.1f} MB > {cfg.ingest.max_file_mb} MB)"
        log.warning("skipping oversized file: %s", path.name, extra={
            "event": "file_too_large", "file": str(path), "size_mb": round(size_mb, 1)})
        return doc

    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        doc.load_error = f"no loader for extension {path.suffix!r}"
        return doc

    pages, error = loader(path)
    doc.pages = pages
    doc.load_error = error

    if error:
        log.warning("could not parse %s: %s", path.name, error,
                    extra={"event": "unparsed_file", "file": str(path), "reason": error})
        return doc

    # A page with almost no characters is a scanned image with no text layer.
    # Report it explicitly rather than silently producing an empty index.
    doc.scanned_pages = [
        p.page_no for p in pages
        if p.char_count < cfg.ingest.scanned_page_char_threshold
    ]
    if doc.scanned_pages and len(doc.scanned_pages) == len(pages):
        doc.load_error = "no extractable text layer (likely a scanned document)"
        log.warning("no text layer in %s -- OCR would be required", path.name,
                    extra={"event": "no_text_layer", "file": str(path)})
    return doc
